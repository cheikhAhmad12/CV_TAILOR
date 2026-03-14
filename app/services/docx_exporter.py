from pathlib import Path
from uuid import uuid4

from docx import Document


def export_resume_to_docx(markdown_text: str) -> str:
    output_dir = Path("/tmp/cv_tailor_exports")
    output_dir.mkdir(parents=True, exist_ok=True)

    filename = f"tailored_resume_{uuid4().hex}.docx"
    output_path = output_dir / filename

    doc = Document()
    for line in markdown_text.splitlines():
        clean_line = line.strip()
        if not clean_line:
            doc.add_paragraph("")
            continue

        if clean_line.startswith("# "):
            doc.add_heading(clean_line[2:].strip(), level=1)
        elif clean_line.startswith("## "):
            doc.add_heading(clean_line[3:].strip(), level=2)
        elif clean_line.startswith("- "):
            doc.add_paragraph(clean_line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(clean_line)

    doc.save(output_path)
    return str(output_path)
